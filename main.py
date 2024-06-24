import streamlit as st
from crewai import Agent, Task, Crew
import os
from crewai_tools import WebScraperTool, SearchTool
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from docx import Document
from io import BytesIO
import base64

# Load environment variables from .env file
load_dotenv()

# Retrieve API keys from environment variables
openai_key = os.getenv("OPENAI_KEY")
serper_key = os.getenv("SERPER_KEY")

# Check if API keys are loaded correctly
if not openai_key:
    raise ValueError("OPENAI_KEY not found. Please set it in the .env file.")
if not serper_key:
    raise ValueError("SERPER_KEY not found. Please set it in the .env file.")

# Set environment variables
os.environ["OPENAI_KEY"] = openai_key
os.environ["SERPER_KEY"] = serper_key

def create_docx(content):
    """Generate a DOCX file with the given content."""
    document = Document()
    document.add_heading('Healthcare Diagnosis and Treatment Recommendations', 0)
    document.add_paragraph(content)
    file_buffer = BytesIO()
    document.save(file_buffer)
    file_buffer.seek(0)
    return file_buffer

def generate_download_link(file_buffer, filename):
    """Generate a download link for the DOCX file."""
    encoded_file = base64.b64encode(file_buffer.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{encoded_file}" download="{filename}">Download Diagnosis and Treatment Plan</a>'

st.set_page_config(layout="wide")

# Title of the app
st.title("Healthcare AI Assistant")

# User inputs for patient information
gender = st.selectbox('Select Gender', ('Male', 'Female', 'Other'))
age = st.number_input('Enter Age', min_value=0, max_value=120, value=25)
symptoms = st.text_area('Enter Symptoms', 'e.g., fever, cough, headache')
medical_history = st.text_area('Enter Medical History', 'e.g., diabetes, hypertension')

# Initialize tools for scraping and searching
scrape_tool = WebScraperTool()
search_tool = SearchTool()

# Initialize language model
language_model = ChatOpenAI(
    model="gpt-3.5-turbo-16k",
    temperature=0.1,
    max_tokens=8000
)

# Define the agents for diagnosis and treatment recommendation
diagnosis_agent = Agent(
    role="Clinical Diagnostician",
    goal="Evaluate patient symptoms and history to determine potential medical conditions.",
    backstory="Specializes in identifying medical conditions by interpreting patient symptoms and historical data, leveraging advanced algorithms and comprehensive medical knowledge.",
    verbose=True,
    allow_delegation=False,
    tools=[search_tool, scrape_tool],
    llm=language_model
)

treatment_agent = Agent(
    role="Medical Treatment Advisor",
    goal="Formulate tailored treatment plans based on diagnostic evaluations.",
    backstory="Focuses on crafting individualized treatment plans by considering diagnoses, patient history, and modern medical practices to recommend optimal therapies.",
    verbose=True,
    allow_delegation=False,
    tools=[search_tool, scrape_tool],
    llm=language_model
)

# Define the tasks for diagnosis and treatment
diagnosis_task = Task(
    description=(
        "1. Review patient symptoms ({symptoms}) and medical history ({medical_history}).\n"
        "2. Offer a preliminary diagnosis with possible conditions based on the input data.\n"
        "3. Restrict the diagnosis to the most probable conditions."
    ),
    expected_output="A preliminary diagnosis outlining potential medical conditions.",
    agent=diagnosis_agent
)

treatment_task = Task(
    description=(
        "1. Develop a step-by-step treatment plan based on the preliminary diagnosis.\n"
        "2. Incorporate the patient's medical history ({medical_history}) and current symptoms ({symptoms}).\n"
        "3. Provide detailed recommendations, including medications, lifestyle changes, and follow-up care instructions."
    ),
    expected_output="A detailed treatment plan customized to the patient's condition.",
    agent=treatment_agent
)

# Create a crew with the defined agents and tasks
medical_crew = Crew(
    agents=[diagnosis_agent, treatment_agent],
    tasks=[diagnosis_task, treatment_task],
    verbose=2
)

# Button to get diagnosis and treatment plan
if st.button("Get Diagnosis and Treatment Plan"):
    with st.spinner('Generating recommendations...'):
        result = medical_crew.kickoff(inputs={"symptoms": symptoms, "medical_history": medical_history})
        st.write(result)
        docx_file = create_docx(result)
        download_link = generate_download_link(docx_file, "diagnosis_and_treatment_plan.docx")
        st.markdown(download_link, unsafe_allow_html=True)
